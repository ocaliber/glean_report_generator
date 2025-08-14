from flask import Flask, redirect, request, session, url_for, send_file, render_template, jsonify
from flask_session import Session
from requests_oauthlib import OAuth2Session
import os
import traceback
import requests
import tempfile
from io import BytesIO
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation

# Word templating
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import json
from collections import defaultdict

# OpenAI for AI summaries
#from openai import OpenAI
import openai
from sp_templates import templates_bp, get_graph_token, SITE_ID

from PIL import Image, ImageDraw, ImageFont



# -----------------------------------------------------------------------------
# Progress-bar generator
# -----------------------------------------------------------------------------
def generate_progress_bar(percentage_used, max_width=300, save_path="progress_bar.png"):
    """
    Generate a professional progress bar with rounded ends and a mask to ensure clean transitions.
    Args:
        percentage_used (float): Percentage of the budget used (0-100 or higher for overspend).
        max_width (int): Maximum width of the bar for 100%.
        save_path (str): Path to save the generated PNG image.
    Returns:
        str: File path to the saved PNG image.
    """
    # Bar dimensions
    bar_height = 10
    radius = bar_height // 2
    text_offset = 15  # Offset for text from the end of the bar
    text_vert_offset = 0.1 # for vertical centralization tweaking.
    dark_teal = (0, 128, 128, 255)
    light_teal = (200, 230, 230, 255)
    light_red = (255, 200, 200, 255)
    dark_red = (229, 57, 53, 255)

    # Calculate the bar width based on overspend
    bar_width = max_width if percentage_used <= 100 else max_width + int((percentage_used - 100) / 100 * max_width)

    # Add percentage text
    font_size = 24
    try:
        font_path = "static/fonts/DejaVuSans-Bold.ttf"  # Replace with your actual font path
        font = ImageFont.truetype(font_path, font_size)
    except OSError:
        print("Custom font not found. Using default font.")
        font = ImageFont.load_default()

    # Adjust text color based on bar color
    text_color = dark_teal[:3] if percentage_used <= 100 else dark_red[:3]
    text = f"{percentage_used:.0f}%"
    text_width, text_height = font.getbbox(text)[2:4]

    # Calculate the overall canvas height (bar height + extra space for tall text)
    canvas_height = max(bar_height, text_height)
    canvas_width = bar_width + text_width + text_offset

    # Center bar and text on the canvas
    bar_y = (canvas_height - bar_height) // 2
    text_y = (canvas_height - text_height) // 2 - text_vert_offset*text_height

    # Create the image with transparent background
    img = Image.new("RGBA", (canvas_width, canvas_height), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    # Create a mask for the obround shape
    mask = Image.new("L", (canvas_width * 2, canvas_height * 2), 0)  # High-res mask for anti-aliasing
    mask_draw = ImageDraw.Draw(mask)
    mask_draw.rounded_rectangle(
        [(0, bar_y * 2), (bar_width * 2, (bar_y + bar_height) * 2)],
        radius=radius * 2,
        fill=255,
    )
    mask = mask.resize((canvas_width, canvas_height), resample=Image.Resampling.LANCZOS)

    # Draw the two segments as rectangles
    progress_position = min(int((percentage_used / 100) * max_width), max_width)
    if percentage_used <= 100:
        # Used portion (dark teal)
        draw.rectangle(
            [(0, bar_y), (progress_position, bar_y + bar_height)],
            fill=dark_teal,
        )
        # Remaining portion (light teal)
        draw.rectangle(
            [(progress_position, bar_y), (max_width, bar_y + bar_height)],
            fill=light_teal,
        )
        # White vertical line at the transition
        draw.line(
            [(progress_position, bar_y), (progress_position, bar_y + bar_height)],
            fill=(255, 255, 255, 255),
            width=3,
        )
    else:
        # Light red for 0 to 100%
        draw.rectangle(
            [(0, bar_y), (max_width, bar_y + bar_height)],
            fill=light_red,
        )
        # Dark red for overspend
        draw.rectangle(
            [(max_width, bar_y), (bar_width, bar_y + bar_height)],
            fill=dark_red,
        )
        # White vertical line at 100%
        draw.line(
            [(max_width, bar_y), (max_width, bar_y + bar_height)],
            fill=(255, 255, 255, 255),
            width=3,
        )

    # Apply the mask to create the obround shape
    img.putalpha(mask)

    # Draw the percentage text next to the bar
    text_x = bar_width + text_offset
    draw.text((text_x, text_y), text, font=font, fill=text_color + (255,))  # Add alpha to text color

    # Save the image
    img.save(save_path, "PNG")
    return save_path

# -----------------------------------------------------------------------------
# Helper functions for nested lists in Word via python-docx
# -----------------------------------------------------------------------------

def find_paragraph_containing_text(document, text):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    return None


def insert_paragraph_before(paragraph, text, style=None):
    """
    Inserts a new paragraph directly before the specified paragraph.
    """
    # create the new paragraph at the same parent
    new_paragraph = paragraph._parent.add_paragraph(text)
    new_el        = new_paragraph._element
    # move it to sit just before the marker paragraph
    paragraph._element.addprevious(new_el)
    if style:
        new_paragraph.style = style
    return new_paragraph

'''
def insert_paragraph_after(paragraph, text, style=None):
    """
    Inserts a new paragraph directly after the specified paragraph in a Word document.
    """
    # Create a new paragraph after the given paragraph in the document structure
    #print(text)
    new_paragraph = paragraph._parent.add_paragraph(text)

    # Move the new paragraph XML element right after the specified paragraph's XML element
    new_paragraph_element = new_paragraph._element
    paragraph._element.addnext(new_paragraph_element)

    new_paragraph.style = style

    return new_paragraph
'''

def set_bullet_level(paragraph, level, numId="1"):
    # Reduced indent per level for tighter, more standard bullet indentation
    indent_per_level = 360
    max_levels = 5
    
    tab_positions = [indent_per_level * (i + 1) for i in range(max_levels)]
    
    if level < 0:
        level = 0
    elif level >= max_levels:
        level = max_levels - 1

    # Get or add the paragraph properties (pPr) for the paragraph
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        numPr = OxmlElement('w:numPr')
        pPr.append(numPr)

    # Set w:ilvl (indentation level) to specify the bullet level
    ilvl = numPr.find(qn('w:ilvl'))
    if ilvl is None:
        ilvl = OxmlElement('w:ilvl')
        numPr.append(ilvl)
    ilvl.set(qn('w:val'), str(level))

    # Set numId to use a standard bullet list
    numId_el = numPr.find(qn('w:numId'))
    if numId_el is None:
        numId_el = OxmlElement('w:numId')
        numPr.append(numId_el)
    numId_el.set(qn('w:val'), numId)  # Using a consistent numId for bullet list

    # Apply "ListBullet" style to ensure bullets instead of numbers
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.append(pStyle)
    pStyle.set(qn('w:val'), "Normal Bulleted List")

    # Adjust tab stops and indentation per level
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)
    for pos in tab_positions[:level + 1]:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'num')
        tab.set(qn('w:pos'), str(pos))
        tabs.append(tab)

    # Add left and hanging indentation per level
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), str(tab_positions[level]))  # Adjust left indent per level
    ind.set(qn('w:hanging'), "360")  # Consistent hanging indent

# Updated insert_bullet_list function with multi-level indentation
def insert_bullet_list(document, summary_text, paragraph, before=False):
    lines = summary_text.strip().split('\n')
    iterable = reversed(lines) if before else lines

    # Process bullet lines from top to bottom
    for line in iterable:
        clean_line = line.lstrip('- ').strip()
        if not clean_line:
            continue  # Skip empty lines

        # Calculate the indent level based on groups of 2 spaces
        leading_spaces = len(line) - len(line.lstrip())
        indent_level = leading_spaces // 2
        #print(f"Raw line: '{line}', Leading spaces: {leading_spaces}, Calculated indent level: {indent_level}")

        # Insert the paragraph with bullet styling after the last bullet paragraph
        bullet_para = insert_paragraph_before(paragraph, clean_line, style='Normal Bulleted List')

        # Use XML manipulation to set the bullet level, tabs, and indentation
        set_bullet_level(bullet_para, indent_level)

        # Update reference paragraph to the last inserted bullet paragraph
        paragraph = bullet_para

def remove_empty_table_rows(doc):
    """
    Walks every table in a python-docx Document and removes any
    rows where *all* cells are empty (i.e. blank or whitespace).
    """
    for table in doc.tables:
        # iterate backwards so removing rows doesn’t shift indices
        for row in list(table.rows)[::-1]:
            if all(cell.text.strip() == "" for cell in row.cells):
                tbl = table._tbl       # the <w:tbl> XML element
                tr  = row._tr          # the <w:tr> XML element
                tbl.remove(tr)

def remove_paragraphs_with_text(document, text_to_remove):
    """
    Removes all paragraphs containing the specified text from a Word document.
    """
    print('Removing pararagraphs containing this text' + text_to_remove)
    # Iterate over each paragraph in the document
    for paragraph in document.paragraphs:
        if text_to_remove in paragraph.text:
            # Remove the paragraph from its parent element
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

    return document

def compute_people_summary(range_entries):
    """
    Returns a list of dicts:
    [
      {
        'name': 'Alice',
        'total': 123.45,
        'projects': [
          {
            'project': 'Project A',
            'total': 80.00,
            'tasks': [
              {'task': 'Design', 'hours': 50.00},
              {'task': 'Review', 'hours': 30.00},
            ]
          },
          {
            'project': 'Project B',
            'total': 43.45,
            'tasks': [ ... ]
          },
        ]
      },
      ...
    ]
    """
    person_map = defaultdict(lambda: {
      'total': Decimal('0.00'),
      'projects': defaultdict(lambda: {
         'total': Decimal('0.00'),
         'tasks': defaultdict(lambda: Decimal('0.00'))
      })
    })

    for e in range_entries:
        if not e.get('billable'):
            continue
        name    = e['user']['name']
        proj    = e['project']['name']
        task    = e['task']['name']
        hrs     = Decimal(str(e['rounded_hours']))

        # accumulate
        person_map[name]['total'] += hrs
        person_map[name]['projects'][proj]['total'] += hrs
        person_map[name]['projects'][proj]['tasks'][task] += hrs

    # build output
    summary = []
    for name, pdata in person_map.items():
        person_total = float(pdata['total'].quantize(Decimal('0.01'), ROUND_HALF_UP))
        projects = []
        for proj, pd in pdata['projects'].items():
            proj_total = float(pd['total'].quantize(Decimal('0.01'), ROUND_HALF_UP))
            tasks = [
              {'task': t, 'hours': float(h.quantize(Decimal('0.01'), ROUND_HALF_UP))}
              for t, h in pd['tasks'].items()
            ]
            projects.append({'project': proj, 'total': proj_total, 'tasks': tasks})
        summary.append({'name': name, 'total': person_total, 'projects': projects})
    return summary


# -----------------------------------------------------------------------------
# Flask & OAuth2 setup
# -----------------------------------------------------------------------------
app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET', 'replace-with-secure-key')
app.config.update({
    'SESSION_COOKIE_SECURE': False,
    'SESSION_COOKIE_HTTPONLY': True,
    'SESSION_COOKIE_SAMESITE': 'Lax',
    'SESSION_TYPE': 'filesystem',
    'SESSION_FILE_DIR': os.path.join(os.getcwd(), 'flask_session'),
})
Session(app)

app.register_blueprint(templates_bp)

# Starting the openai API using azure
openai.api_type    = "azure"
openai.api_base    = os.getenv("AZURE_OPENAI_ENDPOINT")
openai.api_key     = os.getenv("AZURE_OPENAI_API_KEY")
openai.api_version = os.getenv("OPENAI_API_VERSION")

#oa_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
CLIENT_ID     = os.getenv('harvest_client_id')
CLIENT_SECRET = os.getenv('harvest_client_secret')
REDIRECT_URI  = os.getenv('harvest_redirect_uri')
AUTH_URL      = 'https://id.getharvest.com/oauth2/authorize'
TOKEN_URL     = 'https://id.getharvest.com/api/v2/oauth2/token'
API_BASE      = 'https://api.harvestapp.com/v2/'

# -----------------------------------------------------------------------------
# OAuth2 Routes (unchanged)
# -----------------------------------------------------------------------------
@app.route('/')
def index():
    harvest = OAuth2Session(CLIENT_ID, redirect_uri=REDIRECT_URI)
    auth_url, state = harvest.authorization_url(AUTH_URL)
    session['oauth_state'] = state
    return redirect(auth_url)

@app.route('/callback')
def callback():
    # Exchange code for token
    token_data = {
        'client_id': CLIENT_ID,
        'client_secret': CLIENT_SECRET,
        'code': request.args.get('code'),
        'grant_type': 'authorization_code',
        'redirect_uri': REDIRECT_URI,
    }
    resp = OAuth2Session(CLIENT_ID).post(TOKEN_URL, data=token_data)
    session['oauth_token'] = resp.json()
    return redirect(url_for('select_client'))

@app.route('/select_client')
def select_client():
    token = session.get('oauth_token')
    if not token:
        return redirect(url_for('index'))
    harvest = OAuth2Session(CLIENT_ID, token=token)
    try:
        resp_projects = harvest.get(f"{API_BASE}projects?is_active=true")
        resp_projects.raise_for_status()
        projects = resp_projects.json().get('projects', [])
        client_ids = {p['client']['id'] for p in projects}

        resp_clients = harvest.get(f"{API_BASE}clients")
        resp_clients.raise_for_status()
        clients = resp_clients.json().get('clients', [])
        filtered = [c for c in clients if c['id'] in client_ids]
        filtered.sort(key=lambda c: c['name'].lower())

        return render_template('client_selection.html', clients=filtered)
    except Exception:
        traceback.print_exc()
        return "Error fetching clients or projects", 500

@app.route('/submit_selection', methods=['POST'])
def submit_selection():
    client_id = request.form.get('client')
    project_ids = request.form.getlist('project')
    date_range = request.form.get('date_range')
    #template_file = request.files.get('template')
    ai_summarization = request.form.get('ai_summarization', 'no')
    output_format = request.form.get('output_format', 'word')

    if not client_id or not project_ids:
        return "Missing client or project selection", 400

    # 1) Check if the user picked a SharePoint template file
    template_file_id = request.form.get('template_file_id')
    if template_file_id:
        # download that file from Graph
        ms_token = get_graph_token()  # your existing client-credentials helper
        url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drive/items/{template_file_id}/content"
        headers = {"Authorization": f"Bearer {ms_token}"}
        r = requests.get(url, headers=headers)
        r.raise_for_status()
        # save to a temp .docx
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(r.content)
            template_path = tmp.name
    else:
        # fallback to your default .docx
        template_path = 'Report Template.docx'

    start_date, end_date = date_range.split(' - ')
    return redirect(url_for('generate_report',
                            project_ids=','.join(project_ids),
                            start_date=start_date,
                            end_date=end_date,
                            template_path=template_path,
                            ai_summarization=ai_summarization,
                            output_format=output_format))

@app.route('/get_projects/<int:client_id>')
def get_projects(client_id):
    token = session.get('oauth_token')
    if not token:
        return jsonify([]), 401

    harvest = OAuth2Session(CLIENT_ID, token=token)
    try:
        resp = harvest.get(f"{API_BASE}projects?client_id={client_id}&is_active=true")
        resp.raise_for_status()
        projects = resp.json().get('projects', [])
        # return minimal JSON for each project
        return jsonify([{"id": p["id"], "name": p["name"]} for p in projects])
    except Exception:
        traceback.print_exc()
        return jsonify([]), 500

# -----------------------------------------------------------------------------
# Report generation: Word via docxtpl + python-docx
# -----------------------------------------------------------------------------
@app.route('/generate_report/<project_ids>/<start_date>/<end_date>/<path:template_path>')
def generate_report(project_ids, start_date, end_date, template_path):
    # Ensure authenticated
    token = session.get('oauth_token')
    if not token:
        return redirect(url_for('index'))

    harvest = OAuth2Session(CLIENT_ID, token=token)
    tpl = DocxTemplate(template_path)
    
    ids = project_ids.split(',')

    # Fetch data
    projects = []
    all_entries = []
    range_entries = []
    for pid in ids:
        p = harvest.get(f"{API_BASE}projects/{pid}").json()
        projects.append(p)
        all_ = harvest.get(f"{API_BASE}time_entries?project_id={pid}&to={end_date}").json().get('time_entries', [])
        rng_ = harvest.get(f"{API_BASE}time_entries?project_id={pid}&from={start_date}&to={end_date}").json().get('time_entries', [])
        all_entries.extend(all_)
        range_entries.extend(rng_)

    # Prepare context
    reporting_period = f"{start_date} – {end_date}"
    projects_data = []
    for p in projects:
        # 1) Compute and round hours
        rng_h = sum(e['rounded_hours'] for e in range_entries if e['project']['id'] == p['id'])
        tot_h = sum(e['rounded_hours'] for e in all_entries   if e['project']['id'] == p['id'])
        rng_h = round(rng_h, 2)
        tot_h = round(tot_h, 2)

        # 2) Compute and quantize billables to 2 decimal places
        rng_b = sum(
            (
                Decimal(str(e['rounded_hours'])) * Decimal(str(e['billable_rate']))
                for e in range_entries
                if e['project']['id'] == p['id'] and e['billable']
            ),
            Decimal('0.00')
        )
        tot_b = sum(
            (
                Decimal(str(e['rounded_hours'])) * Decimal(str(e['billable_rate']))
                for e in all_entries
                if e['project']['id'] == p['id'] and e['billable']
            ),
            Decimal('0.00')   # <-- ensure the sum starts as a Decimal
        )
        rng_b = rng_b.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        tot_b = tot_b.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        # Convert to plain floats so Jinja won’t try to round them
        rng_b = float(rng_b)
        tot_b = float(tot_b)

        # 3) Budget‐used cell (image or text)
        budget = p.get('budget')
        if budget:
            percent_used = (tot_b / float(budget) * 100)
            img_path = generate_progress_bar(percent_used)
            budget_used = InlineImage(tpl, img_path, width=Inches(1.5))
        else:
            budget_used = "No Budget Assigned"

        # 4) Sum billable rates × hours for any entries not yet billed
        uninvoiced_amt = sum(
            ( Decimal(str(e['rounded_hours'])) * Decimal(str(e['billable_rate'])) )
            for e in range_entries
            if (
                e['project']['id'] == p['id']
                and e.get('billable')        # only true billable entries
                and not e.get('is_billed')   # and not yet invoiced
            )
        ) or Decimal('0.00')
        uninvoiced_amt = uninvoiced_amt.quantize(Decimal('0.01'), ROUND_HALF_UP)
        uninvoiced_amt = float(uninvoiced_amt)

        # 5) getting the total hours set in the project's Team > Monthly Budget
        ## 1) Fetch the project’s per-user budgets
        ua_resp = harvest.get(f"{API_BASE}projects/{p['id']}/user_assignments")
        ua_resp.raise_for_status()
        ua_list = ua_resp.json().get('user_assignments', [])

        ## 2) Sum up all the hourly budgets
        total_monthly_budget_hours = sum(a.get('budget') or 0 for a in ua_list)

        # 6) Calculating Budget
        ##### --- Pull MTD entries once for accurate monthly progress ---
        today = datetime.now().date()
        month_start = today.replace(day=1)
        mtd_entries = harvest.get(
            f"{API_BASE}time_entries?project_id={p['id']}&from={month_start}&to={today}"
        ).json().get('time_entries', [])

        mtd_hours = sum(e.get('rounded_hours', 0) for e in mtd_entries)
        mtd_billable = sum(
            (Decimal(str(e.get('rounded_hours', 0))) * Decimal(str(e.get('billable_rate', 0))))
            for e in mtd_entries if e.get('billable')
        ) or Decimal('0.00')
        mtd_billable = float(mtd_billable.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP))

        # --- Decide how to compute progress ---
        budget_used = "N/A"  # default fallback
        budget_by = (p.get('budget_by') or '').lower()
        proj_budget_val = float(p.get('budget') or 0)
        is_monthly = bool(p.get('budget_is_monthly'))

        def _img(pct):
            img_path = generate_progress_bar(pct)
            return InlineImage(tpl, img_path, width=Inches(1.5))

        try:
            if budget_by in ('project', 'task'):  # hours budget
                if is_monthly or total_monthly_budget_hours:  # monthly roll-over
                    denom = float(total_monthly_budget_hours or proj_budget_val)
                    used  = float(mtd_hours)
                else:  # one-off overall hours budget
                    denom = float(proj_budget_val)
                    used  = float(tot_h)  # cumulative hours to date
                if denom > 0:
                    budget_used = _img(used / denom * 100)

            elif budget_by in ('project_fees', 'task_fees'):  # $ budget
                denom = float(proj_budget_val)
                used  = float(mtd_billable if is_monthly else tot_b)
                if denom > 0:
                    budget_used = _img(used / denom * 100)

            else:
                # No explicit budget type: if you have per-user monthly hours, show MTD vs that
                if total_monthly_budget_hours:
                    denom = float(total_monthly_budget_hours)
                    used  = float(mtd_hours)
                    if denom > 0:
                        budget_used = _img(used / denom * 100)
        except Exception:
            # If anything is off (missing fields, etc.), leave the fallback text
            pass
        #####

        # 7) Building the budget hours etc.             
        budget_hours  = None
        budget_amount = None

        if p.get('budget_by') in ('project', 'task'):
            # hours
            budget_hours = p.get('budget', 0)
        if p.get('budget_by') in ('project_fees', 'task_fees'):
            # dollars
            budget_amount = p.get('budget', 0)

        projects_data.append({
            'code':            p.get('code',''),
            'name':            p.get('name',''),
            'range_hours':     rng_h,
            'total_hours':     tot_h,
            'range_billable':  rng_b,
            'total_billable':  tot_b,
            'budget_used':     budget_used,
            'budget_hours':    budget_hours,
            'budget_amount':   budget_amount,
            'monthly_user_budget_hours': total_monthly_budget_hours,
            'uninvoiced_amount':       uninvoiced_amt,
        })
        
    # Helper to safely sum a numeric key
    def sum_key(list_of_dicts, key):
        total = Decimal('0')
        for d in list_of_dicts:
            val = d.get(key, 0)
            # skip None or empty
            if val is None or (isinstance(val, str) and not val.strip()):
                continue

            # add directly if numeric
            if isinstance(val, (int, float, Decimal)):
                total += Decimal(val)
                continue

            # otherwise try parsing the string
            try:
                total += Decimal(str(val))
            except (InvalidOperation, ValueError):
                # skip anything unparseable
                continue

        # quantize to two decimal places
        total = total.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        return float(total)
        
    project_totals  = {
      'code':            '',
      'name':            'Totals',
      'range_hours':     sum_key(projects_data, 'range_hours'),
      'total_hours':     sum_key(projects_data, 'total_hours'),
      'range_billable':  sum_key(projects_data, 'range_billable'),
      'total_billable':  sum_key(projects_data, 'total_billable'),
      'budget_hours':    sum_key(projects_data, 'budget_hours'),
      'budget_amount':   sum_key(projects_data, 'budget_amount'),
      'monthly_user_budget_hours': sum_key(projects_data, 'monthly_user_budget_hours'),
      'uninvoiced_amount':         sum_key(projects_data, 'uninvoiced_amount'),
    }
    projects_data.append(project_totals)

    # Build entries_data
    entries_data = []
    for e in range_entries:
        entries_data.append({
            'project':       e['project']['name'],
            'project_code':  e['project']['code'],
            'spent_date':    e['spent_date'],
            'task':          e['task']['name'],
            'notes':         e.get('notes',''),
            'rounded_hours': e['rounded_hours'],
            'billable':      e['billable'],
            'user':          {'name': e.get('user',{}).get('name','')},
        })
    
    people_summary = compute_people_summary(range_entries)


    # Render via docxtpl
    context = {
        'reporting_period': reporting_period,
        'projects':         projects_data,
        'entries':          entries_data,
        'people_summary':   people_summary,
    }

    # Save intermediate
    temp_io = BytesIO()
    tpl.render(context)
    tpl.save(temp_io)
    temp_io.seek(0)

    doc = Document(temp_io)

    ## AI Summary    
    AI_START = "[[AI_SUMMARY]]"
    AI_END   = "[[END_AI_SUMMARY]]"

    # Trying to find tags in paragraphs
    start_idx = next(
        (i for i, p in enumerate(tpl.paragraphs)
            if AI_START in p.text),
        None
    )

    if start_idx is not None:
        parts = []

        # 2) Walk forward until we hit a paragraph containing the closing tag
        for p in tpl.paragraphs[start_idx + 1:]:
            if AI_END in p.text:
                break
            parts.append(p.text)
        prompt_template = "\n".join(parts).strip()
    else:
        prompt_template = None

    # If the prompt template exists then we go on to generate the AI summary starting with the json of tasks
    if prompt_template:
        # build the json
        tasks_json = json.dumps(entries_data, indent=2)

        # Construct and call AI summary
        system_prompt = (
        "You are an expert technical writer creating structured summaries "
        "for engineering reports."
        )

        user_prompt = f"{prompt_template}\n\nTasks JSON:\n{tasks_json}"
        
        ai_summary = openai.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt}
            ],
            max_tokens=200,
            temperature=0.5
        ).choices[0].message.content

        print(ai_summary)

        # Find the same paragraph in the rendered doc
        marker = next(
          (p for p in doc.paragraphs
             if p.text.strip().startswith("[[AI_SUMMARY")),
          None
        )
        
        if marker:
            # insert bullets into the rendered doc
            insert_bullet_list(doc, ai_summary, marker, before=True)
            # now strip out the entire prompt block (from [[AI_SUMMARY to ]])
            removing = False
            
            for p in list(doc.paragraphs):
                txt = p.text.strip()
                if txt.startswith("[[AI_SUMMARY"):
                    removing = True
                if removing:
                    p._element.getparent().remove(p._element)
                if txt == "]]":
                    removing = False
        

    # Remove any completely empty table rows that may remain
    from docx.oxml.ns import qn as _qn
    for table in doc.tables:
        # Iterate in reverse to safely remove rows
        for row in list(table.rows)[::-1]:
            if all(cell.text.strip() == '' for cell in row.cells):
                tbl = table._tbl
                tr = row._tr
                tbl.remove(tr)

    remove_empty_table_rows(doc)
    
    fmt = request.args.get('output_format','word').lower()

    out_io = BytesIO()
    doc.save(out_io)
    out_io.seek(0)
    filename = f"Project Report {reporting_period}.docx"

    # Otherwise fall back to Word
    return send_file(
        out_io,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.getenv('PORT',5000)), debug=True)
