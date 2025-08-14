from docx import Document

# Load your Word document
document = Document("Report Template.docx")

def replace_paragraph_text(paragraph, new_text, style=None):
    """
    Replaces the text of the specified paragraph in a Word document.
    """
    # Check if the paragraph is valid
    if paragraph is None:
        raise ValueError("The paragraph provided is None. Make sure the paragraph exists.")

    # Clear the existing text by removing all runs
    for run in paragraph.runs:
        run.text = ""

    # Add a new run with the new text
    new_run = paragraph.add_run(new_text)

    # Apply the style if provided
    if style:
        paragraph.style = style

    return paragraph

# Assume you want to replace text in the 3rd paragraph (index 2)
try:
    paragraph = document.paragraphs[2]  # Adjust the index as needed
except IndexError:
    print("Paragraph index is out of range.")
    paragraph = None

if paragraph:
    replace_paragraph_text(paragraph, "New text here", style=None)
    document.save("modified_example.docx")
else:
    print("The paragraph could not be found.")
